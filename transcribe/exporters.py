"""Multi-format transcript exporters, matches TurboScribe's output set.

Public API:
    export(job, fmt, out_path)            # one call, picks the right encoder
    SUPPORTED_FORMATS                     # ['txt', 'srt', 'vtt', 'csv', 'docx', 'pdf']

Each encoder takes a TranscriptJob and writes to `out_path`. All are pure
Python except DOCX (python-docx, small dep) and PDF (fpdf2, ~1 MB).
"""
from __future__ import annotations

import csv
import logging
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = ('txt', 'srt', 'vtt', 'lrc', 'csv', 'docx', 'pdf')


# ── Time helpers ──────────────────────────────────────────────────────────────

def _fmt_srt_ts(t: float) -> str:
    """SRT timestamp: HH:MM:SS,mmm"""
    if t < 0: t = 0
    h, rem = divmod(t, 3600)
    m, s   = divmod(rem, 60)
    ms     = int((s - int(s)) * 1000)
    return f'{int(h):02d}:{int(m):02d}:{int(s):02d},{ms:03d}'


def _fmt_vtt_ts(t: float) -> str:
    """VTT timestamp: HH:MM:SS.mmm"""
    return _fmt_srt_ts(t).replace(',', '.')


def _fmt_human_ts(t: float) -> str:
    """Reader-friendly timestamp for TXT/DOCX/PDF headings."""
    h, rem = divmod(int(t), 3600)
    m, s   = divmod(rem, 60)
    return f'{h:02d}:{m:02d}:{s:02d}' if h else f'{m:02d}:{s:02d}'


# ── Encoders ──────────────────────────────────────────────────────────────────

def _to_txt(job, out: Path) -> None:
    lines: list[str] = [f'# {Path(job.source).name}']
    if job.summary:
        lines += ['', '## Summary', job.summary, '']
    lines.append('## Transcript')
    cur_speaker = None
    for s in job.segments:
        sp = s.get('speaker') or ''
        ts = _fmt_human_ts(s['start'])
        if sp and sp != cur_speaker:
            lines.append('')
            lines.append(f'[{ts}] {sp}:')
            cur_speaker = sp
        elif not sp:
            lines.append(f'[{ts}]')
        lines.append(s['text'])
    out.write_text('\n'.join(lines), encoding='utf-8')


def _to_srt(job, out: Path) -> None:
    chunks: list[str] = []
    for i, s in enumerate(job.segments, 1):
        sp = s.get('speaker') or ''
        line = f'{sp}: {s["text"]}' if sp else s['text']
        chunks.append(
            f'{i}\n{_fmt_srt_ts(s["start"])} --> {_fmt_srt_ts(s["end"])}\n{line}\n'
        )
    out.write_text('\n'.join(chunks), encoding='utf-8')


def _to_vtt(job, out: Path) -> None:
    chunks: list[str] = ['WEBVTT', '']
    for s in job.segments:
        sp = s.get('speaker') or ''
        line = f'<v {sp}>{s["text"]}</v>' if sp else s['text']
        chunks.append(
            f'{_fmt_vtt_ts(s["start"])} --> {_fmt_vtt_ts(s["end"])}\n{line}\n'
        )
    out.write_text('\n'.join(chunks), encoding='utf-8')


def _fmt_lrc_ts(t: float) -> str:
    """LRC timestamp format: [mm:ss.xx] (centiseconds, not milliseconds)."""
    if t < 0:
        t = 0.0
    m  = int(t // 60)
    s  = t - m * 60
    return f'{m:02d}:{s:05.2f}'


def _to_lrc(job, out: Path) -> None:
    """LRC karaoke / lyric-video format.

    `[mm:ss.xx]line` per segment, supported by Aegisub, VLC's Synchronized
    Lyrics, every karaoke editor, and most lyric-video templates in DaVinci
    Resolve / CapCut / Premiere via simple paste. Optional ID3-style header
    tags (ti / ar / al / length) are emitted only when we have the data so
    we never invent metadata.
    """
    lines: list[str] = []
    meta = getattr(job, 'meta', None) or {}
    title  = meta.get('title')  or getattr(job, 'title',  None)
    artist = meta.get('artist') or getattr(job, 'artist', None)
    album  = meta.get('album')  or getattr(job, 'album',  None)
    # Fallback: derive a clean title from the source filename so karaoke
    # apps and lyric-video templates display something useful instead of
    # an empty header. We strip yt-dlp's trailing "[videoid]" bracket and
    # the file extension to get e.g. "Alors on Joue" from
    # "Alors on Joue [2302333952].m4a".
    if not title and getattr(job, 'source', ''):
        import re as _re, os as _os
        src = job.source
        if src.startswith(('http://', 'https://')):
            base = src.rstrip('/').rsplit('/', 1)[-1]
        else:
            base = _os.path.splitext(_os.path.basename(src))[0]
        cleaned = _re.sub(r'\s*\[[^\]]+\]\s*$', '', base).strip()
        if cleaned:
            title = cleaned
    if title:  lines.append(f'[ti:{title}]')
    if artist: lines.append(f'[ar:{artist}]')
    if album:  lines.append(f'[al:{album}]')
    if job.segments:
        total = job.segments[-1]['end']
        lines.append(f'[length:{_fmt_lrc_ts(total)}]')
    lines.append('')
    for s in job.segments:
        text = (s.get('text') or '').strip()
        if not text:
            continue
        lines.append(f'[{_fmt_lrc_ts(s["start"])}]{text}')
    out.write_text('\n'.join(lines) + '\n', encoding='utf-8')


def _to_csv(job, out: Path) -> None:
    with out.open('w', encoding='utf-8', newline='') as f:
        w = csv.writer(f)
        w.writerow(['start_s', 'end_s', 'speaker', 'text'])
        for s in job.segments:
            w.writerow([
                f'{s["start"]:.3f}',
                f'{s["end"]:.3f}',
                s.get('speaker') or '',
                s['text'],
            ])


def _to_docx(job, out: Path) -> None:
    """DOCX export, speaker headings as bold runs, timestamps as small grey."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        raise RuntimeError('python-docx not installed') from e

    doc = Document()
    doc.add_heading(Path(job.source).name, level=1)
    if job.summary:
        doc.add_heading('Summary', level=2)
        doc.add_paragraph(job.summary)
    doc.add_heading('Transcript', level=2)

    cur_speaker = None
    for s in job.segments:
        sp = s.get('speaker') or ''
        ts = _fmt_human_ts(s['start'])
        p  = doc.add_paragraph()
        if sp and sp != cur_speaker:
            run = p.add_run(f'{sp}  ')
            run.bold = True
            cur_speaker = sp
        ts_run = p.add_run(f'[{ts}] ')
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        p.add_run(s['text'])

    doc.save(str(out))


def _to_pdf(job, out: Path) -> None:
    """PDF export, fpdf2, A4, simple typographic hierarchy."""
    try:
        from fpdf import FPDF
        from fpdf.enums import XPos, YPos
    except ImportError as e:
        raise RuntimeError('fpdf2 not installed') from e

    # The deprecated `ln=True` kwarg leaves the x cursor in a state that
    # makes the next multi_cell() raise "Not enough horizontal space to
    # render a single character". Use the new XPos/YPos API throughout.
    NEXT = dict(new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    # fpdf2's built-in Helvetica handles Latin-1 only; for full Unicode we'd
    # need to embed a TTF. For the MVP, sanitize unsupported codepoints.
    def _safe(t: str) -> str:
        return t.encode('latin-1', 'replace').decode('latin-1')

    pdf.set_font('Helvetica', 'B', 16)
    pdf.cell(0, 10, _safe(Path(job.source).name), **NEXT)
    pdf.ln(2)

    if job.summary:
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, 'Summary', **NEXT)
        pdf.set_font('Helvetica', '', 11)
        pdf.multi_cell(0, 6, _safe(job.summary), **NEXT)
        pdf.ln(4)

    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 8, 'Transcript', **NEXT)
    pdf.ln(2)

    # fpdf2 leaves x at the right margin after multi_cell by default, must
    # pass new_x=LMARGIN explicitly or the next multi_cell raises "Not
    # enough horizontal space to render a single character".
    cur_speaker = None
    for s in job.segments:
        sp = s.get('speaker') or ''
        ts = _fmt_human_ts(s['start'])
        if sp and sp != cur_speaker:
            pdf.ln(2)
            pdf.set_font('Helvetica', 'B', 11)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 6, _safe(sp), **NEXT)
            cur_speaker = sp
        pdf.set_font('Helvetica', '', 11)
        pdf.set_text_color(0, 0, 0)
        pdf.multi_cell(0, 5, _safe(f'[{ts}]  {s["text"]}'), **NEXT)

    pdf.output(str(out))


# ── Dispatcher ────────────────────────────────────────────────────────────────

_ENCODERS: dict[str, Callable] = {
    'txt':  _to_txt,
    'srt':  _to_srt,
    'vtt':  _to_vtt,
    'lrc':  _to_lrc,
    'csv':  _to_csv,
    'docx': _to_docx,
    'pdf':  _to_pdf,
}


def export(job, fmt: str, out_path: str | Path) -> Path:
    """Write `job` to `out_path` in format `fmt`. Returns the resolved path."""
    fmt = fmt.lower().lstrip('.')
    if fmt not in _ENCODERS:
        raise ValueError(f'unsupported format {fmt!r}, pick from {SUPPORTED_FORMATS}')
    out = Path(out_path).expanduser().resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    _ENCODERS[fmt](job, out)
    logger.info(f'exported {fmt} → {out}')
    return out
